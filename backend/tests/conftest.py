import pytest
import asyncio
import logging
import os
import sys
import tempfile
import shutil
from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# addding backend/app path for import
sys.path.insert(0, str(Path(__file__).parent.parent))

@pytest.fixture(scope="session")
def event_loop():
    """Fixture pour gérer l'event loop asyncio"""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()

@pytest.fixture
def test_logger():
    """Logger pour les tests"""
    logger = logging.getLogger('test_logger')
    logger.setLevel(logging.DEBUG)
    
    handler = logging.StreamHandler()
    handler.setLevel(logging.DEBUG)
    formatter = logging.Formatter('%(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    
    if not logger.handlers:
        logger.addHandler(handler)
    
    return logger

@pytest.fixture
def mock_config():
    """Configuration mock pour les tests"""
    return {
        'google_sheet_url': 'https://test-url',
        'api_timeout': 30
    }


@pytest.fixture(scope="session")
def test_files_dir():
    test_dir = tempfile.mkdtemp()
    yield test_dir
    # cleaning after all tests
    shutil.rmtree(test_dir, ignore_errors=True)

@pytest.fixture(scope="session")
def sample_files(test_files_dir):
    files = {}

    pdf_path = os.path.join(test_files_dir, "valid_test.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(100, 750, "Ceci est un texte de test dans un PDF réel")
    c.drawString(100, 730, "Ceci est la deuxième ligne.")
    c.showPage()
    c.save()

    files['valid_pdf'] = {
        'path': pdf_path,
        'expected_text': "Ceci est un texte de test dans un PDF réel.\n Ceci est la deuxièmé ligne"
    }

    # pdf multi-pages
    multipages_path = os.path.join(test_files_dir, "multipages_test.pdf")
    page_contents = ["Page 1 contenu", "Page 2 contenu", "Page 3 contenu"]
    c = canvas.Canvas(multipages_path, pagesize=letter)

    for content in page_contents:
        c.drawString(100, 750, content)
        c.showPage()
    c.save()

    files['multipages_pdf'] = {
        'path': multipages_path,
        'pages_contents': page_contents
    }

    # empty pdf
    empty_pdf_path = os.path.join(test_files_dir, "empty_test.pdf")
    c = canvas.Canvas(empty_pdf_path, pagesize=letter)
    c.showPage()
    c.save()
    files['empty_pdf'] = empty_pdf_path

    # docx
    docx_path = os.path.join(test_files_dir, "valid_test.docx")
    doc = Document()
    doc.add_paragraph("Ceci est un texte de test dans un docx réel.")
    doc.add_paragraph("Avec plusieurs lignes de contenu.")
    doc.save(docx_path)

    files['valid_docx'] = {
        'path': docx_path,
        'expected_text': "Ceci est un texte de test dans un docx réel.\nAvec plusieurs lignes de contenu."
    }

    # empty docx
    empty_docx_path = os.path.join(test_files_dir, "empty_test.docx")
    doc = Document()
    doc.save(empty_docx_path)
    files['empty_docx'] = empty_docx_path

    # corrupted files
    corrupted_pdf_path = os.path.join(test_files_dir, "corrupted.pdf")
    with open(corrupted_pdf_path, "wb") as f:
        f.write(b"Ceci n'est pas un vrai PDF")

    files['corrupted_pdf'] = corrupted_pdf_path

    corrupted_docx_path = os.path.join(test_files_dir, "corrupted.docx")
    with open(corrupted_docx_path, "wb") as f:
        f.write(b"Ceci n'est pas un vrai DOCX")

    files['corrupted_docx'] = corrupted_docx_path

    return files




