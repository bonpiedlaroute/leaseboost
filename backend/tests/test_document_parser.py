import pytest
import os
from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter



from app.services.document_parser import DocumentParser


@pytest.mark.asyncio
async def test_extract_extract_from_real_pdf_success(sample_files):

    parser = DocumentParser()

    with open(sample_files['valid_pdf']['path'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_pdf(file_content)

    assert result is not None
    assert isinstance(result, str)
    assert "Ceci est un texte de test" in result
    assert "deuxième ligne" in result

@pytest.mark.asyncio
async def test_extract_from_real_pdf_multipages(sample_files):

    parser = DocumentParser()

    with open(sample_files['multipages_pdf']['path'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_pdf(file_content)

    assert result is not None

    for content in sample_files['multipages_pdf']['pages_contents']:
        assert content in result
    
    lines = result.split('\n')
    assert len(lines) >= 3

@pytest.mark.asyncio
async def test_extract_from_real_empty_pdf(sample_files):

    parser = DocumentParser()

    with open(sample_files['empty_pdf'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_pdf(file_content)

    # pdf with no content should return None
    assert result is None

@pytest.mark.asyncio
async def test_extract_from_corrupted_pdf(sample_files):

    parser = DocumentParser()

    with open(sample_files['corrupted_pdf'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_pdf(file_content)

    assert result is None

@pytest.mark.asyncio
async def test_large_pdf_handling(test_files_dir):

    parser = DocumentParser()

    pdf_path = os.path.join(test_files_dir, "large_test.pdf")
    
    c = canvas.Canvas(pdf_path, pagesize=letter)

    for page_num in range(50):
        c.drawString(100, 750, f"Contenu de la page {page_num + 1}")
        c.showPage()

    c.save()
    try:
        with open(pdf_path, 'rb') as file:
            file_content = file.read()

        result = parser._extract_from_pdf(file_content)

        assert result is not None
        assert "page 1" in result
        assert "page 50" in result
    finally:
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)

@pytest.mark.asyncio
async def test_extract_from_real_docx_success(sample_files):

    parser = DocumentParser()

    with open(sample_files['valid_docx']['path'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_docx(file_content)

    assert result is not None
    assert isinstance(result, str)
    assert "Ceci est un texte de test" in result
    assert "plusieurs lignes" in result


@pytest.mark.asyncio
async def test_extract_from_real_doc_format(sample_files):
    parser = DocumentParser()

    with open(sample_files['valid_docx']['path'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_docx(file_content)

    assert result is not None
    assert isinstance(result, str)
    assert "Ceci est un texte de test" in result
 
@pytest.mark.asyncio
async def test_extract_from_corrupted_docs(sample_files):

    parser = DocumentParser()

    with open(sample_files['corrupted_docx'], 'rb') as file:
        file_content = file.read()

    result = parser._extract_from_docx(file_content)

    assert result is None

@pytest.mark.asyncio
async def test_validate_inputs_empty_content():
    parser = DocumentParser()

    result = await parser.extract_text_from_file(b"", "test.pdf")

    assert result is None

@pytest.mark.asyncio
async def test_validate_inputs_invalid_filename():
    parser = DocumentParser()

    result = await parser.extract_text_from_file(b"content", "filename_without_extension")

    assert result is None


@pytest.mark.asyncio
async def test_unsupported_file_format():

    parser = DocumentParser()

    result = await parser.extract_text_from_file(b"content", "test.txt")

    assert result is None


@pytest.mark.asyncio
async def test_pdf_with_special_characters(test_files_dir):
    
    pdf_path = os.path.join(test_files_dir, "special_chars.pdf")
    
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(100, 750, "Texte avec caractères spéciaux: àéèùç €")
    c.showPage()
    c.save()
    
    parser= DocumentParser()
    
    with open(pdf_path, 'rb') as f:
        file_content = f.read()
    
    result = await parser.extract_text_from_file(file_content, "special_chars.pdf")
    
    assert result is not None
    assert "Texte avec" in result


@pytest.mark.asyncio
async def test_docx_with_formatting(test_files_dir):
    
    docx_path = os.path.join(test_files_dir, "formatted.docx")
    
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Texte normal. ")
    p.add_run("Texte en gras.").bold = True
    p.add_run(" Texte en italique.").italic = True
    doc.save(docx_path)
    
    parser = DocumentParser()
    
    with open(docx_path, 'rb') as f:
        file_content = f.read()
    
    result = await parser.extract_text_from_file(file_content, "formatted.docx")
    
    assert result is not None
    assert "Texte normal" in result
    assert "Texte en gras" in result
    assert "Texte en italique" in result

if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])